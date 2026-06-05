#!/usr/bin/env python3
"""
Genera curso-demo-rcp.docx · Word de ejemplo con el formato ideal
para entregar el contenido de un curso al template OVA Waygroup.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Márgenes
for s in doc.sections:
    s.top_margin = Inches(0.8)
    s.bottom_margin = Inches(0.8)
    s.left_margin = Inches(0.9)
    s.right_margin = Inches(0.9)

# Estilos rápidos
def h1(t):
    p = doc.add_heading(t, level=1)
    for r in p.runs:
        r.font.color.rgb = RGBColor(0xF3, 0x70, 0x21)
    return p

def h2(t):
    return doc.add_heading(t, level=2)

def p(t, bold=False, italic=False, color=None, size=None):
    para = doc.add_paragraph()
    run = para.add_run(t)
    run.bold = bold
    run.italic = italic
    if color: run.font.color.rgb = color
    if size: run.font.size = Pt(size)
    return para

def hint(t):
    para = doc.add_paragraph()
    run = para.add_run(f"[hint: {t}]")
    run.italic = True
    run.font.color.rgb = RGBColor(0xCA, 0x8A, 0x04)
    run.font.size = Pt(10)

def label_value(lbl, val):
    para = doc.add_paragraph()
    r1 = para.add_run(f"{lbl}: ")
    r1.bold = True
    para.add_run(val)

def divider():
    para = doc.add_paragraph()
    run = para.add_run("─" * 40)
    run.font.color.rgb = RGBColor(0xB7, 0xC3, 0xD4)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

def bullet(t, level=0):
    para = doc.add_paragraph(t, style='List Bullet' if level == 0 else 'List Bullet 2')

# ════════════════════════════════════════════════════════════
# CABECERA DEL CURSO
# ════════════════════════════════════════════════════════════
h1("CURSO · Metadatos generales")
label_value("Nombre", "RCP Básico · Entrenamiento Inmersivo")
label_value("Subtítulo", "Programa educativo basado en guías AHA 2025")
label_value("Código", "RCP-2025")
label_value("Cliente", "Waygroup For Education LATAM")
label_value("Audiencia", "Primer respondiente (lego)")
label_value("Duración", "20–30 min · incluye debriefing")
label_value("Objetivo", "Capacitar al primer respondiente en la identificación del paro "
                        "cardiorrespiratorio y la ejecución del protocolo RCP básico con DEA.")

h2("Preset visual")
label_value("Preset", "vital  (1 de los 14 disponibles · ver presets.html)")
label_value("Versión default", "hibrida  (hibrida | clara | oscura)")
label_value("Logo", "adjunto-aparte.svg")

h2("Estructura del curso")
estructura = [
    "01 · Portada",
    "02 · Bienvenida",
    "03 · Pregunta detonante",
    "04 · Conceptos previos",
    "05 · Infarto vs paro",
    "06 · Cadena de supervivencia",
    "07 · RCP de calidad",
    "08 · Uso del DEA",
    "09 · Después del paro",
    "10 · Las 8 fases del entrenamiento",
    "11 · El simulador en acción",
    "12 · Evaluación",
    "13 · Glosario",
    "14 · Referencias y créditos",
]
for item in estructura:
    bullet(item)

divider()

# ════════════════════════════════════════════════════════════
# PANTALLA 01 · PORTADA
# ════════════════════════════════════════════════════════════
h1("PANTALLA 01 · Portada")
hint("hero portada default")

label_value("Eyebrow", "CURSO RCP-2025 · ISO 9001:2015 · AHA 2025")
label_value("Título", "Reanimación Cardiopulmonar")
label_value("Énfasis (palabras en color)", "en VR háptica")
label_value("Subtítulo", "Programa educativo inmersivo Waygroup. El conocimiento "
                         "se aprende leyendo. La competencia se construye actuando.")

p("Stats inline:", bold=True)
bullet("20–30 min | Duración")
bullet("8 fases | Modelo pedagógico")
bullet("VR + Háptica | Modalidad")
bullet("Primer respondiente | Audiencia")

label_value("CTA principal", "Comenzar el curso")
label_value("CTA secundario", "Descargar material PDF")

divider()

# ════════════════════════════════════════════════════════════
# PANTALLA 02 · BIENVENIDA
# ════════════════════════════════════════════════════════════
h1("PANTALLA 02 · Bienvenida")
hint("ficha técnica + grid 2 con objetivos")

label_value("Título", "Te damos la bienvenida al programa")
label_value("Lead", "Antes de entrar al simulador háptico, te presentamos el modelo "
                    "de educación inmersiva y lo que vas a vivir en las próximas pantallas.")

p("Ficha técnica (4 ítems):", bold=True)
bullet("Nombre | RCP Básica · Entrenamiento Inmersivo")
bullet("Duración | 20–30 min · incluye debriefing")
bullet("Objetivo | Identificar paro y ejecutar protocolo RCP + DEA")
bullet("Audiencia | Primer respondiente no profesional")

p("Resultados de aprendizaje (lado derecho del grid):", bold=True)
bullet("Identificar signos de paro cardiorrespiratorio en menos de 10 segundos")
bullet("Ejecutar compresiones torácicas de alta calidad (5–6 cm, 100–120/min)")
bullet("Operar un DEA siguiendo el protocolo seguro")
bullet("Activar la cadena de supervivencia con delegación específica")

divider()

# ════════════════════════════════════════════════════════════
# PANTALLA 03 · PREGUNTA DETONANTE
# ════════════════════════════════════════════════════════════
h1("PANTALLA 03 · Pregunta detonante")
hint("bloque detonante grande + 4 concept cards con link")

label_value("Título", "¿Qué pasa cuando un corazón se detiene?")
label_value("Lead", "Antes de responder, observá este dato. Define todo lo que harás "
                    "en las próximas pantallas.")

p("Dato impactante (bloque dark grande):", bold=True)
bullet("Valor en grande: 7–10%")
bullet("Texto: Cada minuto sin RCP reduce la probabilidad de sobrevida entre un "
       "7% y un 10%. El daño cerebral irreversible comienza en 4 a 6 minutos.")
bullet("Fuente: American Heart Association · Guías AHA 2025")

label_value("Puente (texto entre stat y cards)",
            "Para responder esta pregunta necesitás dominar 4 conceptos clave.")

p("Concept cards (4 con link a la pantalla siguiente):", bold=True)
bullet("01 · Paro cardíaco → conceptos.html · Por qué se detiene el corazón.")
bullet("02 · Infarto vs paro → comparativo.html · La diferencia clínica clave.")
bullet("03 · Cadena de supervivencia → cadena.html · Los 6 eslabones de AHA 2025.")
bullet("04 · RCP de calidad → compresiones.html · Profundidad, frecuencia, continuidad.")

divider()

# ════════════════════════════════════════════════════════════
# PANTALLA 04 · CONCEPTOS PREVIOS
# ════════════════════════════════════════════════════════════
h1("PANTALLA 04 · Conceptos previos")
hint("acordeón con 5 definiciones + callout final con novedad AHA 2025")

label_value("Eyebrow", "Pantalla 04 · Definiciones clave")
label_value("Título", "Conceptos previos")
label_value("Lead", "Andamiaje cognitivo antes de la simulación. Tocá cada "
                    "definición para expandirla.")

p("Definiciones (ítems del acordeón):", bold=True)

p("1. Paro o parada cardíaca", bold=True)
p("Se produce cuando el corazón deja de latir o presenta un ritmo irregular "
  "rápido que impide el bombeo efectivo de sangre. Sin intervención, el daño "
  "cerebral irreversible comienza en 4–6 minutos.")

p("2. Reanimación Cardiopulmonar (RCP)", bold=True)
p("Conjunto de competencias para mantener la circulación hasta que llegue el SEM. "
  "Para primeros respondientes, AHA 2025 recomienda Hands-Only CPR: compresiones "
  "continuas sin respiraciones boca a boca.")

p("3. Desfibrilador Externo Automático (DEA)", bold=True)
p("Dispositivo electrónico que analiza el ritmo cardíaco y guía al usuario "
  "para administrar una descarga si detecta un ritmo desfibrilable. "
  "La desfibrilación en los primeros 3–5 minutos eleva la sobrevida del 5–10% al 50–70%.")

p("4. Verificación de pulso", bold=True)
p("AHA 2025 reafirma: NO se recomienda verificar pulso por parte de primeros "
  "respondientes (sensibilidad menor al 50%). Si la víctima no responde y no "
  "respira normalmente, asumir paro e iniciar compresiones de inmediato.")

p("5. Servicio de Emergencias Médicas (SEM)", bold=True)
p("Sistema profesional de respuesta extrahospitalaria. En Colombia se activa "
  "marcando el 123.")

p("Callout final:", bold=True)
bullet("Tipo: warning  (default | success | warning | danger)")
bullet("Tag: Novedad AHA 2025")
bullet("Título: Cambio terminológico importante")
bullet('Texto: El término "respiraciones de rescate" se elimina. Ahora se usa '
       '"respiraciones" para asistencia durante RCP, y "ventilación" se reserva '
       "para soporte mecánico por profesionales.")

divider()

# ════════════════════════════════════════════════════════════
# PANTALLA 05 · COMPARATIVO
# ════════════════════════════════════════════════════════════
h1("PANTALLA 05 · Infarto vs Paro")
hint("comparativo VS · 2 cards lado a lado")

label_value("Título", "Infarto cardíaco vs Paro cardíaco")
label_value("Lead", "Son condiciones distintas que requieren respuestas distintas. "
                    "Confundirlas cuesta tiempo crítico.")

p("Lado A:", bold=True)
bullet("Tag: Infarto cardíaco")
bullet("Título: Problema de circulación")
bullet("Descripción: El suministro de sangre al músculo cardíaco se reduce u obstruye.")
p("Bullets:")
bullet("Arteria coronaria obstruida", level=1)
bullet("El corazón sigue latiendo", level=1)
bullet("Persona usualmente consciente", level=1)
bullet("Síntomas: dolor torácico, sudoración, náusea", level=1)
bullet("Si no se trata, puede derivar en paro", level=1)

p("Lado B:", bold=True)
bullet("Tag: Paro cardíaco")
bullet("Título: Problema eléctrico")
bullet("Descripción: La función eléctrica del corazón se detiene súbitamente. "
       "No hay bombeo efectivo.")
p("Bullets:")
bullet("El corazón deja de latir", level=1)
bullet("Pérdida inmediata de respuesta", level=1)
bullet("No respira o respira anormalmente", level=1)
bullet("Requiere RCP + DEA ahora", level=1)
bullet("Daño cerebral en 4–6 minutos", level=1)

p("Callout final:", bold=True)
bullet("Tipo: warning")
bullet("Tag: Regla del primer respondiente")
bullet("Título: Ante la duda, actuá")
bullet("Texto: Si la víctima no responde y no respira normalmente, asumí paro "
       "e iniciá compresiones. AHA 2025 reafirma que las compresiones imperfectas "
       "siempre son mejores que la ausencia de compresiones.")

divider()

# ════════════════════════════════════════════════════════════
# PANTALLA 06 · CADENA
# ════════════════════════════════════════════════════════════
h1("PANTALLA 06 · Cadena de supervivencia")
hint("hero dark con cadena de 6 eslabones + bullets detallados por eslabón")

label_value("Eyebrow", "Pantalla 06 · AHA 2025")
label_value("Título", "Cadena de supervivencia unificada")
label_value("Lead", "Las guías AHA 2025 reemplazan las cuatro cadenas separadas "
                    "de 2020 por una sola cadena de seis eslabones.")

p("Cadena de 6 eslabones (componente cadena horizontal):", bold=True)
bullet("01 · Reconocimiento — Primer respondiente")
bullet("02 · Activar SEM — Primer respondiente")
bullet("03 · RCP de calidad — Primer respondiente")
bullet("04 · Desfibrilación — DEA + respondiente")
bullet("05 · Posresucitación — SEM")
bullet("06 · Recuperación — Hospital")

p("Detalle de los 3 eslabones del primer respondiente (uno por bloque):", bold=True)

p("Eslabón 1: Reconocimiento y activación", bold=True)
bullet("Asegurá la escena · evaluá riesgos mirando arriba, abajo, lados, frente y atrás.")
bullet("Verificá respuesta · tocá los hombros y llamá al paciente.")
bullet("NO verifiques pulso · AHA 2025 reafirma esto.")
bullet("Llamá al 123 · da ubicación, edad y estado.")
bullet("Delegá específicamente · señalá a una persona concreta.")

p("Eslabón 2: RCP de alta calidad", bold=True)
bullet("Posicionate · costado del paciente, rodillas en piso, espalda y codos rectos.")
bullet("Ubicá el punto de compresión · apófisis xifoides + 3 dedos hacia arriba.")
bullet("Coloca las manos · una encima de la otra, tenar sobre el esternón.")
bullet("Comprimí fuerte y rápido · 5–6 cm, 100–120/min, descompresión completa.")
bullet("No te detengas · pausas menores a 10 segundos.")

p("Eslabón 3: Desfibrilación temprana", bold=True)
bullet("Encendé el DEA y seguí las instrucciones de voz.")
bullet("Colocá los parches en el tórax desnudo y seco.")
bullet("Alejate · verificá que nadie toque al paciente durante la descarga.")
bullet("Persona del DEA debe ser distinta a la de las compresiones.")
bullet("Continuá hasta que llegue la ambulancia o el paciente reaccione.")

divider()

# ════════════════════════════════════════════════════════════
# PANTALLA 07 · COMPRESIONES
# ════════════════════════════════════════════════════════════
h1("PANTALLA 07 · RCP de calidad")
hint("3 KPIs arriba + 3 figuras del posicionamiento + tabla 3-col 'interfaz simulador' + callout musical")

label_value("Eyebrow", "Pantalla 07 · Eslabón 3 · Compresiones")
label_value("Título", "Compresiones de alta calidad")
label_value("Lead", "Tres parámetros se miden en tiempo real con el dummy "
                    "háptico Bluetooth y la interfaz tipo reloj del simulador.")

p("KPIs (3):", bold=True)
bullet("Profundidad | 5–6 cm | ~1/3 del diámetro del tórax. Permitir descompresión completa.")
bullet("Frecuencia | 100–120 | compresiones por minuto. Tempo de canciones tipo Stayin' Alive.")
bullet("Fracción | >60% | del tiempo total. Pausas <10 seg. Iniciar lo antes posible.")

p("Posicionamiento del primer respondiente (3 figuras):", bold=True)
bullet("Figura 1: Posición correcta (perfil arrodillado con callouts 1-2-3-4)")
bullet("Figura 2: Referencia anatómica (xifoides + 3 dedos)")
bullet("Figura 3: Reloj de retroalimentación (dial dark · variante --dark)")

p("Interfaz de retroalimentación del simulador (tabla 3 columnas):", bold=True)
bullet("Profundidad | 5–6 cm | Números laterales cambian de naranja a azul con la profundidad correcta. Permitir descompresión completa.")
bullet("Velocidad | 100–120/min | Aguja central en zona verde. Fuera del rango = sangre no circula o calidad se pierde.")
bullet("Inicio rápido | <10 seg | Empezar lo antes posible. Asignar a alguien la llamada y el DEA.")

p("Callout final:", bold=True)
bullet("Tipo: default (brand)")
bullet("Tag: Tip clínico")
bullet("Título: El truco de las canciones")
bullet("Texto: Cantar mentalmente Stayin' Alive, La Macarena o Baby Shark "
       "mantiene exactamente el tempo de 100–120 cpm. AHA lo recomienda oficialmente.")

divider()

# ════════════════════════════════════════════════════════════
# PANTALLA 08 · USO DEL DEA
# ════════════════════════════════════════════════════════════
h1("PANTALLA 08 · Uso del DEA")
hint("stat impactante (×7) + stepper vertical 6 pasos + 2 callouts (coordinación + novedad mujeres)")

label_value("Eyebrow", "Pantalla 08 · Eslabón 4 · Desfibrilación")
label_value("Título", "Uso del Desfibrilador Externo Automático")
label_value("Lead", "La desfibrilación dentro de los primeros 3 a 5 minutos eleva "
                    "la sobrevida del 5-10% al 50-70%.")

p("Stat destacado:", bold=True)
bullet("Valor: ×7")
bullet("Texto: Multiplica por 7 las posibilidades de salvar la vida cuando se "
       "aplica en los primeros minutos.")

p("Stepper · 6 pasos secuenciales:", bold=True)
bullet("01 · Activar — Encendé el DEA · Abrir tapa o presionar power.")
bullet("02 · Preparar — Tórax desnudo y seco · Cortar ropa, secar.")
bullet("03 · Colocar — Pegar los parches según la imagen impresa.")
bullet("04 · ¡Importante! — Alejarse · Cuando el DEA diga 'no toque'.")
bullet("05 · Seguir — Obedecer las indicaciones de voz.")
bullet("06 · Continuar — Reanudar compresiones en <5 segundos tras la descarga.")

p("Callout 1:", bold=True)
bullet("Tipo: default · Tag: Coordinación")
bullet("Título: DEA y compresiones · personas distintas")
bullet("Texto: Quien opera el DEA debe ser distinta a quien hace las compresiones.")

p("Callout 2:", bold=True)
bullet("Tipo: warning · Tag: Novedad AHA 2025 · DEA en mujeres")
bullet("Título: Ajustar el sostén, no retirarlo")
bullet("Texto: Para reducir la hesitación, AHA 2025 establece que es razonable "
       "ajustar el sostén (moverlo arriba o abajo) en lugar de retirarlo completamente.")

divider()

# ════════════════════════════════════════════════════════════
# PANTALLA 12 · EVALUACIÓN
# ════════════════════════════════════════════════════════════
h1("PANTALLA 12 · Evaluación")
hint("quiz con 6 preguntas + card final 'estás listo para el simulador'")

label_value("Título", "Verifica tu comprensión")
label_value("Lead", "Antes de ingresar al simulador háptico, validá los conceptos clave.")

p("Pregunta 1:", bold=True)
bullet("Texto: Según AHA 2025, ¿qué debe hacer un primer respondiente ante una "
       "víctima que no responde y no respira normalmente?")
bullet("Opción A: Verificar pulso carotídeo durante 10 segundos.")
bullet("Opción B: Asumir paro cardíaco e iniciar compresiones inmediatamente. ★ CORRECTA")
bullet("Opción C: Esperar a que llegue el personal médico.")
bullet("Opción D: Realizar 2 respiraciones de rescate antes de comprimir.")
bullet("Feedback correcto: ✓ AHA 2025 elimina la verificación de pulso por parte "
       "de legos (sensibilidad <50%). Cada minuto sin RCP reduce la sobrevida 7-10%.")
bullet("Feedback incorrecto: ✗ Revisá la pantalla 04. Si no responde y no "
       "respira normal, asumir paro y comprimir.")

p("Pregunta 2:", bold=True)
bullet("Texto: ¿Cuál es la profundidad correcta de las compresiones torácicas en un adulto?")
bullet("Opción A: 3 a 4 cm")
bullet("Opción B: 5 a 6 cm ★ CORRECTA")
bullet("Opción C: 7 a 8 cm")
bullet("Opción D: Lo más profundo posible sin medir")

p("[... preguntas 3-6 con el mismo formato ...]", italic=True)

p("Card final post-quiz:", bold=True)
bullet("Icono: trophy")
bullet("Título: Estás listo para el simulador")
bullet("CTA: Conocer el simulador → simulador.html")

divider()

# ════════════════════════════════════════════════════════════
# GLOSARIO
# ════════════════════════════════════════════════════════════
h1("PANTALLA 13 · Glosario")
hint("auto-renderizado · entregar como lista plana letra-término-definición")

p("Términos del glosario:", bold=True)
bullet("A | AHA | American Heart Association · publica las guías de RCP.")
bullet("A | Apófisis xifoides | Prominencia ósea en el extremo inferior del esternón.")
bullet("C | Cadena de supervivencia | Secuencia de 6 eslabones AHA 2025.")
bullet("D | DEA | Desfibrilador Externo Automático.")
bullet("F | Fibrilación ventricular | Ritmo cardíaco caótico no efectivo.")
bullet("H | Hands-Only CPR | RCP solo con las manos · recomendada por AHA 2025.")
bullet("L | Lego | Persona sin formación profesional sanitaria.")
bullet("P | Primer respondiente | Persona que actúa primero ante una emergencia.")
bullet("S | SEM | Servicio de Emergencias Médicas · en Colombia se activa con 123.")
bullet("T | Tenar | Zona carnosa de la palma cerca del pulgar.")

divider()

# ════════════════════════════════════════════════════════════
# REFERENCIAS Y CRÉDITOS
# ════════════════════════════════════════════════════════════
h1("PANTALLA 14 · Referencias y créditos")
hint("auto-renderizado · lista plana")

p("Referencias bibliográficas:", bold=True)
bullet("American Heart Association (2025). Guías AHA 2025 para RCP y "
       "Atención Cardiovascular de Emergencia.")
bullet("International Liaison Committee on Resuscitation (2025). "
       "Consenso internacional ILCOR CoSTR 2025.")
bullet("Otras referencias relevantes del curso...")

p("Créditos del equipo:", bold=True)
bullet("Bloque 1: Equipo Waygroup For Education LATAM")
bullet("  - Dirección académica · Nombre Apellido")
bullet("  - Diseño instruccional · Nombre Apellido")
bullet("  - Desarrollo VR · Nombre Apellido")
bullet("  - Diseño UX/UI · Nombre Apellido")
bullet("  - Validación clínica · Nombre Apellido")

label_value("Certificación", "ISO 9001:2015")
label_value("Licencia", "Creative Commons BY-NC-SA 4.0")

divider()

# ════════════════════════════════════════════════════════════
# NOTAS FINALES
# ════════════════════════════════════════════════════════════
h1("Notas finales para la diseñadora instruccional")
p("Este documento es un EJEMPLO del formato ideal. Lo importante:", bold=True)
bullet("Usá Heading 1 para cada PANTALLA")
bullet("Usá Heading 2 para subsecciones")
bullet("Marcá el [hint: componente] al inicio de cada pantalla (opcional pero "
       "súper recomendado)")
bullet("Separá pantallas con una línea o salto de página")
bullet("Listas con viñetas para enumeraciones")
bullet("Negrita para etiquetas/labels")
bullet("Si no sabés qué componente usar, no marqués hint y yo lo decido por vos")

p("Cualquier duda, preguntame en chat antes de armar el Word completo.", italic=True)

# Guardar
import sys
out_path = sys.argv[1] if len(sys.argv) > 1 else "curso-demo-rcp.docx"
doc.save(out_path)
print(f"✓ Generado: {out_path}")
